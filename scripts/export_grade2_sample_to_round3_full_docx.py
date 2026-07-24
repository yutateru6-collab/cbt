import json
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = ROOT / "exports" / "grade2-sample-to-round3-word"
SOURCE_FILES = [
    "grade2-set-01.js",
    "grade2-vocab-sets.js",
    "grade2-listening-part2-sets.js",
    "exam-data.js",
]
TARGET_KEYS = ["sample", "set-01", "set-02", "set-03"]

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(91, 102, 116)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F8"


def load_sets():
    node_script = r"""
const fs = require('fs');
const vm = require('vm');
const context = { window: {} };
vm.createContext(context);
for (const file of ['grade2-set-01.js','grade2-vocab-sets.js','grade2-listening-part2-sets.js','exam-data.js']) {
  vm.runInContext(fs.readFileSync(file, 'utf8'), context, { filename: file });
}
const targetKeys = ['sample', 'set-01', 'set-02', 'set-03'];
const sets = context.window.examData.grades.grade2.sets
  .filter(set => targetKeys.includes(set.key || set.setId))
  .map(set => ({
    key: set.key || set.setId,
    label: set.label,
    description: set.description || '',
    readingPages: set.readingPages || [],
    listeningQuestions: set.listeningQuestions || [],
  }))
  .sort((a, b) => targetKeys.indexOf(a.key) - targetKeys.indexOf(b.key));
process.stdout.write(JSON.stringify(sets));
"""
    completed = subprocess.run(
        ["node", "-e", node_script],
        cwd=ROOT,
        check=True,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )
    return json.loads(completed.stdout)


def validate_sets(exam_sets):
    if [exam_set["key"] for exam_set in exam_sets] != TARGET_KEYS:
        raise RuntimeError(
            f"Expected sets {TARGET_KEYS}, found {[exam_set['key'] for exam_set in exam_sets]}"
        )

    for exam_set in exam_sets:
        reading_pages = exam_set.get("readingPages") or []
        reading_questions = [
            question
            for page in reading_pages
            for question in (page.get("questions") or [])
        ]
        listening_questions = exam_set.get("listeningQuestions") or []

        if len(reading_questions) != 31:
            raise RuntimeError(
                f"{exam_set['label']}: expected 31 reading questions, found {len(reading_questions)}"
            )
        if len(listening_questions) != 30:
            raise RuntimeError(
                f"{exam_set['label']}: expected 30 listening questions, found {len(listening_questions)}"
            )

        for question in reading_questions:
            validate_question(exam_set["label"], "reading", question, "text")
        for question in listening_questions:
            validate_question(exam_set["label"], "listening", question, "questionText")
            if not str(question.get("script", "")).strip():
                raise RuntimeError(
                    f"{exam_set['label']} listening No.{question.get('id')}: script missing"
                )


def validate_question(set_label, category, question, text_key):
    question_id = question.get("id")
    if not str(question.get(text_key, "")).strip():
        raise RuntimeError(f"{set_label} {category} No.{question_id}: {text_key} missing")
    choices = question.get("choices") or []
    if len(choices) != 4:
        raise RuntimeError(
            f"{set_label} {category} No.{question_id}: expected 4 choices, found {len(choices)}"
        )
    correct = question.get("correct")
    if not isinstance(correct, int) or not 1 <= correct <= 4:
        raise RuntimeError(
            f"{set_label} {category} No.{question_id}: invalid correct choice {correct}"
        )
    if not str(question.get("explanation", "")).strip():
        raise RuntimeError(
            f"{set_label} {category} No.{question_id}: explanation missing"
        )


def set_style_font(style, latin_name="Calibri", east_asia_name="Yu Gothic", size=11):
    style.font.name = latin_name
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:ascii"), latin_name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), latin_name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_name)


def set_run_font(
    run,
    size=11,
    bold=False,
    italic=False,
    color=INK,
    latin_name="Calibri",
    east_asia_name="Yu Gothic",
):
    run.font.name = latin_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia_name)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade_paragraph(paragraph, fill):
    paragraph_properties = paragraph._p.get_or_add_pPr()
    shading = paragraph_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        paragraph_properties.append(shading)
    shading.set(qn("w:fill"), fill)


def add_field(run, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction_text, separate, value, end])


def configure_document(running_title):
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    set_style_font(normal, size=11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading1 = document.styles["Heading 1"]
    set_style_font(heading1, size=16)
    heading1.font.bold = True
    heading1.font.color.rgb = BLUE
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(10)
    heading1.paragraph_format.keep_with_next = True

    heading2 = document.styles["Heading 2"]
    set_style_font(heading2, size=13)
    heading2.font.bold = True
    heading2.font.color.rgb = DARK_BLUE
    heading2.paragraph_format.space_before = Pt(14)
    heading2.paragraph_format.space_after = Pt(7)
    heading2.paragraph_format.keep_with_next = True

    heading3 = document.styles["Heading 3"]
    set_style_font(heading3, size=12)
    heading3.font.bold = True
    heading3.font.color.rgb = DARK_BLUE
    heading3.paragraph_format.space_before = Pt(10)
    heading3.paragraph_format.space_after = Pt(5)
    heading3.paragraph_format.keep_with_next = True

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(header_paragraph, after=0, line=1.0)
    set_run_font(
        header_paragraph.add_run(running_title),
        size=8.5,
        color=MUTED,
    )

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(footer_paragraph, after=0, line=1.0)
    set_run_font(footer_paragraph.add_run("—  "), size=8.5, color=MUTED)
    page_run = footer_paragraph.add_run()
    set_run_font(page_run, size=8.5, color=MUTED)
    add_field(page_run, "PAGE")
    set_run_font(footer_paragraph.add_run("  —"), size=8.5, color=MUTED)

    return document


def add_cover(document, title, subtitle, scope_lines):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=82, after=12, line=1.0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(paragraph.add_run("英検2級 模擬問題"), size=12, bold=True, color=BLUE)

    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=12, line=1.12)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(paragraph.add_run(title), size=25, bold=True, color=INK)

    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=40, line=1.1)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(paragraph.add_run(subtitle), size=12, color=MUTED)

    for line in scope_lines:
        paragraph = document.add_paragraph()
        set_paragraph_spacing(paragraph, after=6, line=1.0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(paragraph.add_run(line), size=10.5, color=DARK_BLUE)

    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=52, after=0, line=1.0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        paragraph.add_run("問題・リスニングスクリプト・正解・解説を収録"),
        size=9.5,
        color=MUTED,
    )
    document.add_page_break()


def add_set_overview(document, exam_set):
    document.add_heading(exam_set["label"], level=1)

    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=12, line=1.25)
    shade_paragraph(paragraph, LIGHT_BLUE)
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.right_indent = Inches(0.15)
    set_run_font(
        paragraph.add_run(
            "収録内容：リーディング31問／リスニング30問（Part 1・Part 2）／全問の正解・解説"
        ),
        size=10.5,
        bold=True,
        color=DARK_BLUE,
    )

    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=8)
    set_run_font(
        paragraph.add_run(
            "問題編のあとに解答・解説編を配置しています。リスニング問題には、読み上げ原稿、Question、選択肢を収録しています。"
        ),
        size=10.5,
    )
    document.add_page_break()


def add_instruction(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.right_indent = Inches(0.15)
    set_paragraph_spacing(paragraph, before=0, after=8, line=1.2)
    shade_paragraph(paragraph, LIGHT_BLUE)
    set_run_font(paragraph.add_run(str(text).strip()), size=10.5, bold=True, color=DARK_BLUE)


def add_question_choices(document, question, text_key="text"):
    heading = document.add_paragraph()
    heading.paragraph_format.keep_with_next = True
    set_paragraph_spacing(heading, before=8, after=3, line=1.0)
    set_run_font(heading.add_run(f"No. {question['id']}"), size=11.5, bold=True, color=DARK_BLUE)

    question_paragraph = document.add_paragraph()
    question_paragraph.paragraph_format.keep_with_next = True
    set_paragraph_spacing(question_paragraph, after=4, line=1.22)
    set_run_font(question_paragraph.add_run(str(question[text_key]).strip()), size=11)

    choices = question["choices"]
    for choice_index, choice in enumerate(choices, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.38)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = choice_index < len(choices)
        set_paragraph_spacing(paragraph, after=2, line=1.15)
        set_run_font(
            paragraph.add_run(f"({choice_index})  {choice}"),
            size=10.5,
        )


def add_reading_problem_section(document, exam_set):
    document.add_heading("問題編｜リーディング", level=1)

    for page in exam_set["readingPages"]:
        document.add_heading(page["label"], level=2)
        if page.get("instruction"):
            add_instruction(document, page["instruction"])

        passage_title = str(page.get("passageTitle", "")).strip()
        passage = page.get("passage") or []
        if passage_title:
            title_paragraph = document.add_paragraph()
            title_paragraph.paragraph_format.keep_with_next = True
            set_paragraph_spacing(title_paragraph, before=4, after=5, line=1.0)
            set_run_font(title_paragraph.add_run(passage_title), size=12, bold=True, color=INK)
        for passage_paragraph_text in passage:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.12)
            paragraph.paragraph_format.right_indent = Inches(0.12)
            paragraph.paragraph_format.keep_together = True
            set_paragraph_spacing(paragraph, after=6, line=1.3)
            set_run_font(paragraph.add_run(str(passage_paragraph_text).strip()), size=10.7)

        if passage:
            separator = document.add_paragraph()
            set_paragraph_spacing(separator, after=2, line=1.0)
            shade_paragraph(separator, LIGHT_GRAY)
            set_run_font(separator.add_run("設問"), size=9.5, bold=True, color=MUTED)

        for question in page.get("questions") or []:
            add_question_choices(document, question)


def split_speakers(script):
    normalized = re.sub(r"\s+([AB]):\s*", r"\n\1: ", str(script).strip())
    return normalized.lstrip()


def add_script(document, script):
    label = document.add_paragraph()
    label.paragraph_format.keep_with_next = True
    set_paragraph_spacing(label, before=1, after=2, line=1.0)
    set_run_font(label.add_run("読み上げ原稿"), size=9.5, bold=True, color=MUTED)

    script_paragraph = document.add_paragraph()
    script_paragraph.paragraph_format.left_indent = Inches(0.12)
    script_paragraph.paragraph_format.right_indent = Inches(0.12)
    script_paragraph.paragraph_format.keep_together = True
    set_paragraph_spacing(script_paragraph, after=5, line=1.22)
    shade_paragraph(script_paragraph, LIGHT_GRAY)
    lines = split_speakers(script).splitlines()
    for line_index, line in enumerate(lines):
        if line_index:
            script_paragraph.add_run().add_break(WD_BREAK.LINE)
        speaker_match = re.match(r"^([AB]):\s*(.*)$", line)
        if speaker_match:
            set_run_font(
                script_paragraph.add_run(f"{speaker_match.group(1)}: "),
                size=10.5,
                bold=True,
                color=DARK_BLUE,
            )
            set_run_font(script_paragraph.add_run(speaker_match.group(2)), size=10.5)
        else:
            set_run_font(script_paragraph.add_run(line), size=10.5)


def add_listening_problem_section(document, exam_set):
    document.add_page_break()
    document.add_heading("問題編｜リスニング", level=1)

    current_part = None
    for question in exam_set["listeningQuestions"]:
        part = question.get("part") or "Listening"
        if part != current_part:
            document.add_heading(part, level=2)
            current_part = part

        heading = document.add_paragraph()
        heading.paragraph_format.keep_with_next = True
        set_paragraph_spacing(heading, before=9, after=3, line=1.0)
        set_run_font(
            heading.add_run(f"No. {question['id']}"),
            size=11.5,
            bold=True,
            color=DARK_BLUE,
        )
        add_script(document, question["script"])

        question_label = document.add_paragraph()
        question_label.paragraph_format.keep_with_next = True
        set_paragraph_spacing(question_label, after=2, line=1.0)
        set_run_font(question_label.add_run("Question"), size=9.5, bold=True, color=MUTED)

        question_paragraph = document.add_paragraph()
        question_paragraph.paragraph_format.keep_with_next = True
        set_paragraph_spacing(question_paragraph, after=4, line=1.18)
        set_run_font(
            question_paragraph.add_run(str(question["questionText"]).strip()),
            size=10.8,
            bold=True,
        )

        choices = question["choices"]
        for choice_index, choice in enumerate(choices, start=1):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.38)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.keep_with_next = choice_index < len(choices)
            set_paragraph_spacing(paragraph, after=2, line=1.12)
            set_run_font(paragraph.add_run(f"({choice_index})  {choice}"), size=10.4)


def flatten_reading_questions(exam_set):
    return [
        question
        for page in exam_set["readingPages"]
        for question in (page.get("questions") or [])
    ]


def add_answer_index(document, questions):
    for row_start in range(0, len(questions), 5):
        row_questions = questions[row_start : row_start + 5]
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_together = True
        set_paragraph_spacing(paragraph, after=4, line=1.1)
        text = "　｜　".join(
            f"No.{question['id']} ({question['correct']})" for question in row_questions
        )
        set_run_font(paragraph.add_run(text), size=10.3, bold=True, color=DARK_BLUE)


def add_answer_block(document, question):
    correct_index = question["correct"]
    correct_choice = question["choices"][correct_index - 1]

    heading = document.add_paragraph()
    heading.paragraph_format.keep_with_next = True
    set_paragraph_spacing(heading, before=8, after=2, line=1.0)
    set_run_font(
        heading.add_run(f"No. {question['id']}　正解 ({correct_index}) {correct_choice}"),
        size=10.8,
        bold=True,
        color=DARK_BLUE,
    )

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.keep_together = True
    set_paragraph_spacing(paragraph, after=5, line=1.22)
    set_run_font(paragraph.add_run(str(question["explanation"]).strip()), size=10.4)


def add_answer_section(document, exam_set):
    document.add_page_break()
    document.add_heading("解答・解説編", level=1)

    reading_questions = flatten_reading_questions(exam_set)
    listening_questions = exam_set["listeningQuestions"]

    document.add_heading("リーディング｜正解一覧", level=2)
    add_answer_index(document, reading_questions)
    document.add_heading("リーディング｜解説", level=2)
    for question in reading_questions:
        add_answer_block(document, question)

    document.add_page_break()
    document.add_heading("リスニング｜正解一覧", level=2)
    add_answer_index(document, listening_questions)
    document.add_heading("リスニング｜解説", level=2)
    for question in listening_questions:
        add_answer_block(document, question)


def add_set(document, exam_set, add_leading_break=False):
    if add_leading_break:
        document.add_page_break()
    add_set_overview(document, exam_set)
    add_reading_problem_section(document, exam_set)
    add_listening_problem_section(document, exam_set)
    add_answer_section(document, exam_set)


def safe_filename_label(label):
    return re.sub(r'[<>:"/\\|?*]', "_", label)


def save_document(document, output_path, title, subject):
    document.core_properties.title = title
    document.core_properties.subject = subject
    document.core_properties.author = "CBT形式4技能トレーニング"
    document.core_properties.keywords = "英検2級, リーディング, リスニング, スクリプト, 解答, 解説"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def build_single_set_document(exam_set):
    title = f"2級 {exam_set['label']}｜リーディング・リスニング"
    document = configure_document(title)
    add_cover(
        document,
        f"{exam_set['label']} 完全収録",
        "リーディング・リスニング・スクリプト・解答解説",
        ["リーディング 31問", "リスニング 30問（Part 1・Part 2）"],
    )
    add_set(document, exam_set)

    filename = (
        f"2級_{safe_filename_label(exam_set['label'])}"
        "_リーディング・リスニング・スクリプト・解答解説.docx"
    )
    output_path = OUTPUT_DIR / filename
    save_document(
        document,
        output_path,
        f"2級 {exam_set['label']} リーディング・リスニング・スクリプト・解答解説",
        f"{exam_set['label']}の全61問と解答解説",
    )
    return output_path


def build_combined_document(exam_sets):
    title = "2級 サンプル〜第3回｜総合版"
    document = configure_document(title)
    add_cover(
        document,
        "サンプル〜第3回 総合版",
        "リーディング・リスニング・スクリプト・解答解説",
        [
            "サンプル問題・第1回・第2回・第3回",
            "リーディング 全124問",
            "リスニング 全120問（Part 1・Part 2）",
        ],
    )

    for index, exam_set in enumerate(exam_sets):
        add_set(document, exam_set, add_leading_break=index > 0)

    output_path = (
        OUTPUT_DIR
        / "2級_サンプル-第3回_リーディング・リスニング・スクリプト・解答解説_総合版.docx"
    )
    save_document(
        document,
        output_path,
        "2級 サンプル〜第3回 リーディング・リスニング・スクリプト・解答解説 総合版",
        "サンプル問題から第3回までの全244問、リスニング原稿、正解、解説",
    )
    return output_path


def main():
    exam_sets = load_sets()
    validate_sets(exam_sets)

    output_paths = [build_single_set_document(exam_set) for exam_set in exam_sets]
    output_paths.append(build_combined_document(exam_sets))

    print(f"source_files={','.join(SOURCE_FILES)}")
    print(f"sets={len(exam_sets)}")
    print("reading_questions=124")
    print("listening_questions=120")
    for output_path in output_paths:
        print(output_path)


if __name__ == "__main__":
    main()

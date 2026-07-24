import json
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = ROOT / "exports"
OUTPUT_PATH = OUTPUT_DIR / "2級_リスニング_第1回-第5回_全スクリプトと選択肢.docx"


def load_grade2_listening_sets():
    node_script = r"""
const fs = require('fs');
const vm = require('vm');
const context = { window: {} };
vm.createContext(context);
for (const file of ['grade2-set-01.js','grade2-vocab-sets.js','grade2-listening-part2-sets.js','exam-data.js']) {
  vm.runInContext(fs.readFileSync(file, 'utf8'), context, { filename: file });
}
const sets = context.window.examData.grades.grade2.sets
  .filter(set => /^set-0[1-5]$/.test(set.key))
  .map(set => ({ key: set.key, label: set.label, listeningQuestions: set.listeningQuestions }));
process.stdout.write(JSON.stringify(sets));
"""
    completed = subprocess.run(
        ["node", "-e", node_script],
        cwd=ROOT,
        check=True,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )
    return json.loads(completed.stdout)


def validate_sets(sets):
    if len(sets) != 5:
        raise RuntimeError(f"Expected 5 sets, found {len(sets)}")

    questions = []
    for exam_set in sets:
        set_questions = exam_set.get("listeningQuestions") or []
        if len(set_questions) != 30:
            raise RuntimeError(f"{exam_set['label']}: expected 30 questions, found {len(set_questions)}")
        for question in set_questions:
            if not str(question.get("script", "")).strip():
                raise RuntimeError(f"{exam_set['label']} No.{question.get('id')}: script missing")
            if not str(question.get("questionText", "")).strip():
                raise RuntimeError(f"{exam_set['label']} No.{question.get('id')}: question text missing")
            if len(question.get("choices") or []) != 4:
                raise RuntimeError(f"{exam_set['label']} No.{question.get('id')}: choices missing")
            questions.append((exam_set, question))

    if len(questions) != 150:
        raise RuntimeError(f"Expected 150 questions, found {len(questions)}")
    return questions


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_spacing(paragraph, before=0, after=4, line=1.15):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_label(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    set_paragraph_spacing(paragraph, before=3, after=2, line=1.0)
    run = paragraph.add_run(text)
    set_run_font(run, size=10, bold=True, color=(64, 74, 80))


def split_speakers(script):
    return re.sub(r"\s+([AB]):\s*", r"\n\1: ", script.strip()).lstrip()


def configure_document():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15
    return document


def add_title(document):
    paragraph = document.add_paragraph()
    paragraph.alignment = 1
    set_paragraph_spacing(paragraph, after=4, line=1.0)
    set_run_font(paragraph.add_run("2級 リスニング 第1回〜第5回"), size=18, bold=True, color=(17, 24, 39))

    paragraph = document.add_paragraph()
    paragraph.alignment = 1
    set_paragraph_spacing(paragraph, after=14, line=1.0)
    set_run_font(paragraph.add_run("全150問　読み上げ原稿・質問・選択肢"), size=11, color=(93, 100, 112))


def add_question(document, exam_set, question, index):
    heading = document.add_paragraph()
    heading.paragraph_format.keep_with_next = True
    set_paragraph_spacing(heading, before=12 if index > 1 else 0, after=5, line=1.0)
    part = question.get("part") or question.get("section") or ""
    set_run_font(
        heading.add_run(f"{exam_set['label']}  No. {question['id']}  ({part})"),
        size=13,
        bold=True,
        color=(17, 24, 39),
    )

    add_label(document, "原稿")
    script_paragraph = document.add_paragraph()
    script_paragraph.paragraph_format.keep_together = True
    set_paragraph_spacing(script_paragraph, after=6, line=1.2)
    for line_index, line in enumerate(split_speakers(str(question["script"])).splitlines()):
        if line_index:
            script_paragraph.add_run().add_break(WD_BREAK.LINE)
        speaker_match = re.match(r"^([AB]):\s*(.*)$", line)
        if speaker_match:
            set_run_font(script_paragraph.add_run(f"{speaker_match.group(1)}: "), size=10.5, bold=True)
            set_run_font(script_paragraph.add_run(speaker_match.group(2)), size=10.5)
        else:
            set_run_font(script_paragraph.add_run(line), size=10.5)

    add_label(document, "質問")
    question_paragraph = document.add_paragraph()
    question_paragraph.paragraph_format.keep_together = True
    set_paragraph_spacing(question_paragraph, after=6, line=1.15)
    set_run_font(question_paragraph.add_run(str(question["questionText"]).strip()), size=10.5, bold=True)

    add_label(document, "選択肢")
    choices_table = document.add_table(rows=0, cols=2)
    choices_table.autofit = False
    choices_table.columns[0].width = Cm(0.9)
    choices_table.columns[1].width = Cm(15.2)
    for choice_index, choice in enumerate(question["choices"], start=1):
        number_cell, choice_cell = choices_table.add_row().cells
        number_cell.width = Cm(0.9)
        choice_cell.width = Cm(15.2)
        set_cell_shading(number_cell, "EEF1F4")

        number_paragraph = number_cell.paragraphs[0]
        number_paragraph.alignment = 1
        set_paragraph_spacing(number_paragraph, after=0, line=1.0)
        set_run_font(number_paragraph.add_run(str(choice_index)), size=10, bold=True, color=(64, 74, 80))

        choice_paragraph = choice_cell.paragraphs[0]
        set_paragraph_spacing(choice_paragraph, after=0, line=1.0)
        set_run_font(choice_paragraph.add_run(str(choice)), size=10.5)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def main():
    sets = load_grade2_listening_sets()
    questions = validate_sets(sets)
    document = configure_document()
    add_title(document)
    for index, (exam_set, question) in enumerate(questions, start=1):
        add_question(document, exam_set, question, index)

    document.core_properties.title = "2級 リスニング 第1回〜第5回 全スクリプトと選択肢"
    document.core_properties.subject = "2級リスニング全150問の読み上げ原稿、質問文、選択肢"
    document.core_properties.author = "CBT形式4技能トレーニング"
    OUTPUT_DIR.mkdir(exist_ok=True)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)
    print(f"sets={len(sets)} questions={len(questions)}")


if __name__ == "__main__":
    main()

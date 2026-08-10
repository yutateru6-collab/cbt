"""Create the Grade 2 reading vocabulary reference for Sets 1-3.

The source questions and passages are loaded through Node exactly as the app
loads them.  The output follows the ``compact_reference_guide`` DOCX preset.
"""

from __future__ import annotations

import json
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = ROOT / "exports" / "英検2級_第1回-第3回_リーディング重要語彙リスト.docx"
SOURCE_FILES = [
    "grade2-set-01.js",
    "grade2-vocab-sets.js",
    "grade2-speaking-sets.js",
    "grade2-listening-part2-sets.js",
    "grade2-set-01-explanations.js",
    "grade2-skill-explanations.js",
    "exam-data.js",
]
SET_KEYS = ("set-01", "set-02", "set-03")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(91, 102, 116)
HEADER_FILL = "E8EEF5"


CHOICE_LEXICON = {
    "permit": ("パーミット", "許可証"),
    "grant": ("グラント", "助成金・補助金"),
    "receipt": ("レシート", "領収書"),
    "survey": ("サーベイ", "調査"),
    "weighed": ("ウェイド", "重さを量った"),
    "labeled": ("レイベルド", "ラベルを付けた"),
    "sealed": ("シールド", "封をした"),
    "inspected": ("インスペクティッド", "詳しく点検した"),
    "current": ("カレント", "現在の・最新の"),
    "basic": ("ベイシック", "基本的な"),
    "local": ("ローカル", "地元の"),
    "public": ("パブリック", "公の・公共の"),
    "softly": ("ソフトリー", "やわらかく"),
    "warmly": ("ウォームリー", "温かく"),
    "firmly": ("ファームリー", "きっぱりと・強く"),
    "casually": ("カジュアリー", "気軽に"),
    "permission": ("パーミッション", "許可"),
    "evidence": ("エヴィデンス", "証拠"),
    "equipment": ("イクイップメント", "設備・道具"),
    "attention": ("アテンション", "注意"),
    "reserve": ("リザーヴ", "予約する・確保する"),
    "check": ("チェック", "確認する"),
    "change": ("チェインジ", "変更する"),
    "choose": ("チューズ", "選ぶ"),
    "covered": ("カヴァード", "覆われた・屋根のある"),
    "crowded": ("クラウディッド", "混雑した"),
    "regular": ("レギュラー", "通常の"),
    "temporary": ("テンポラリー", "一時的な"),
    "method": ("メソッド", "方法"),
    "request": ("リクエスト", "依頼"),
    "priority": ("プライオリティ", "優先事項"),
    "tradition": ("トラディション", "伝統"),
    "persuade": ("パースウェイド", "説得する"),
    "warn": ("ウォーン", "警告する"),
    "remind": ("リマインド", "思い出させる"),
    "inform": ("インフォーム", "知らせる"),
    "slightly": ("スライトリー", "わずかに"),
    "recently": ("リーセントリー", "最近"),
    "severely": ("セヴィアリー", "ひどく・深刻に"),
    "deliberately": ("デリバレットリー", "故意に"),
    "plan for": ("プラン フォー", "〜に備えて計画する"),
    "keep track of": ("キープ トラック オヴ", "〜を記録して把握する"),
    "cut down on": ("カット ダウン オン", "〜を減らす"),
    "look back on": ("ルック バック オン", "〜を振り返る"),
    "turn in": ("ターン イン", "提出する"),
    "send out": ("センド アウト", "発送する・配布する"),
    "carry out": ("キャリー アウト", "実施する"),
    "look over": ("ルック オウヴァー", "ざっと確認する"),
    "account for": ("アカウント フォー", "〜を説明する・占める"),
    "care for": ("ケア フォー", "〜の世話をする"),
    "stand for": ("スタンド フォー", "〜を表す"),
    "call for": ("コール フォー", "〜を必要とする"),
    "go through": ("ゴウ スルー", "よく確認する"),
    "draw up": ("ドロー アップ", "作成する"),
    "hand in": ("ハンド イン", "提出する"),
    "add to": ("アッド トゥ", "〜に加える"),
    "as long as": ("アズ ロング アズ", "〜する限り・〜という条件で"),
    "even though": ("イーヴン ゾウ", "〜だけれども"),
    "whether or not": ("ウェザー オア ノット", "〜かどうかにかかわらず"),
    "unless": ("アンレス", "〜でない限り"),
    "no sooner": ("ノウ スーナー", "〜するとすぐに"),
    "no less": ("ノウ レス", "それでもなお・まさに"),
    "no longer": ("ノウ ロンガー", "もはや〜ではない"),
    "no better": ("ノウ ベター", "少しも良くない"),
    "at walking distance": ("アット ウォーキング ディスタンス", "歩ける距離で（不自然な形）"),
    "in walking distance": ("イン ウォーキング ディスタンス", "歩ける距離で（不自然な形）"),
    "to walking distance": ("トゥ ウォーキング ディスタンス", "歩ける距離へ（不自然な形）"),
    "within walking distance": ("ウィズイン ウォーキング ディスタンス", "歩いて行ける距離に"),
    "description": ("ディスクリプション", "説明・記述"),
    "replacement": ("リプレイスメント", "代わりの品"),
    "direction": ("ダイレクション", "方向・指示"),
    "durable": ("デュラブル", "丈夫な・耐久性のある"),
    "powerful": ("パワフル", "強力な"),
    "expensive": ("イクスペンシヴ", "高価な"),
    "portable": ("ポータブル", "持ち運びできる"),
    "declined": ("ディクラインド", "断った"),
    "accepted": ("アクセプティッド", "受け入れた"),
    "suggested": ("サジェスティッド", "提案した"),
    "considered": ("コンシダード", "検討した"),
    "silently": ("サイレントリー", "静かに"),
    "accurately": ("アキュラットリー", "正確に"),
    "calmly": ("カームリー", "落ち着いて"),
    "appetite": ("アペタイト", "食欲"),
    "interest": ("インタレスト", "興味"),
    "patience": ("ペイシェンス", "忍耐"),
    "energy": ("エナジー", "体力・エネルギー"),
    "inspect": ("インスペクト", "詳しく調べる"),
    "label": ("レイベル", "ラベルを付ける"),
    "store": ("ストア", "保管する"),
    "identify": ("アイデンティファイ", "見分ける・特定する"),
    "ordinary": ("オーディナリー", "普通の"),
    "efficient": ("イフィシェント", "効率のよい"),
    "private": ("プライヴェット", "私的な・専用の"),
    "familiar": ("ファミリア", "よく知られた・なじみのある"),
    "besides": ("ビサイズ", "そのうえ"),
    "otherwise": ("アザーワイズ", "そうでなければ"),
    "instead": ("インステッド", "代わりに"),
    "meanwhile": ("ミーンワイル", "その間に"),
    "habit": ("ハビット", "習慣"),
    "benefit": ("ベネフィット", "利点"),
    "obstacle": ("オブスタクル", "障害・妨げ"),
    "sample": ("サンプル", "見本"),
    "predict": ("プリディクト", "予測する"),
    "estimate": ("エスティメイト", "見積もる・推定する"),
    "verify": ("ヴェリファイ", "正しいか確認する"),
    "revise": ("リヴァイズ", "修正する"),
    "cut back on": ("カット バック オン", "〜を減らす"),
    "catch up on": ("キャッチ アップ オン", "遅れを取り戻す"),
    "stand up for": ("スタンド アップ フォー", "〜を擁護する"),
    "look out for": ("ルック アウト フォー", "〜に気をつける"),
    "fill in for": ("フィル イン フォー", "〜の代わりをする"),
    "go along with": ("ゴウ アロング ウィズ", "〜に同意する・従う"),
    "register for": ("レジスター フォー", "〜に登録する"),
    "unsubscribe from": ("アンサブスクライブ フロム", "〜の購読・登録を解除する"),
    "upgrade to": ("アップグレイド トゥ", "〜へ上位変更する"),
    "pay for": ("ペイ フォー", "〜の料金を払う"),
    "keep an eye on": ("キープ アン アイ オン", "〜を注意して見守る"),
    "make use of": ("メイク ユース オヴ", "〜を利用する"),
    "get rid of": ("ゲット リッド オヴ", "〜を取り除く"),
    "catch sight of": ("キャッチ サイト オヴ", "〜をちらっと見る"),
    "in exchange for": ("イン イクスチェインジ フォー", "〜と引き換えに"),
    "in contrast to": ("イン コントラスト トゥ", "〜とは対照的に"),
    "in preparation for": ("イン プレパレイション フォー", "〜に備えて"),
    "in addition to": ("イン アディション トゥ", "〜に加えて"),
    "out of place": ("アウト オヴ プレイス", "場違いな"),
    "out of date": ("アウト オヴ デイト", "時代遅れの"),
    "out of order": ("アウト オヴ オーダー", "故障中の"),
    "out of work": ("アウト オヴ ワーク", "失業中の"),
    "on average": ("オン アヴェレッジ", "平均して"),
    "in public": ("イン パブリック", "人前で"),
    "by mistake": ("バイ ミステイク", "誤って"),
    "in advance": ("イン アドヴァンス", "前もって"),
    "restriction": ("リストリクション", "制限"),
    "decision": ("ディシジョン", "決定"),
    "exception": ("イクセプション", "例外"),
    "locate": ("ロウケイト", "位置を特定する"),
    "transport": ("トランスポート", "運ぶ"),
    "examine": ("イグザミン", "詳しく調べる"),
    "replace": ("リプレイス", "取り替える"),
    "detailed": ("ディテイルド", "詳しい"),
    "popular": ("ポピュラー", "人気がある"),
    "convenient": ("コンヴィーニエント", "便利な"),
    "reliable": ("リライアブル", "信頼できる"),
    "discount": ("ディスカウント", "値引き"),
    "refund": ("リファンド", "返金"),
    "deposit": ("デポジット", "預け金・保証金"),
    "bill": ("ビル", "請求書・請求額"),
    "suddenly": ("サドゥンリー", "突然"),
    "patiently": ("ペイシェントリー", "辛抱強く"),
    "quietly": ("クワイエットリー", "静かに"),
    "jokingly": ("ジョウキングリー", "冗談っぽく"),
    "compare": ("コンペア", "比較する"),
    "describe": ("ディスクライブ", "説明する"),
    "consider": ("コンシダー", "検討する"),
    "adopt": ("アドプト", "採用する・導入する"),
    "permanent": ("パーマネント", "恒久的な"),
    "recommendation": ("レコメンデイション", "勧め・推薦"),
    "advantage": ("アドヴァンテージ", "利点"),
    "requirement": ("リクワイアメント", "必要条件"),
    "choice": ("チョイス", "選択"),
    "deliver": ("デリヴァー", "届ける・発表する"),
    "gather": ("ギャザー", "集める"),
    "gradually": ("グラジュアリー", "徐々に"),
    "briefly": ("ブリーフリー", "短時間に・簡潔に"),
    "locally": ("ローカリー", "地元で・局所的に"),
    "secretly": ("シークレットリー", "ひそかに"),
    "looked after": ("ルックト アフター", "世話をした"),
    "ran out of": ("ラン アウト オヴ", "〜を使い切った"),
    "kept track of": ("ケプト トラック オヴ", "〜を記録して把握した"),
    "put up with": ("プット アップ ウィズ", "〜を我慢する"),
    "come up with": ("カム アップ ウィズ", "〜を思いつく"),
    "keep up with": ("キープ アップ ウィズ", "〜に遅れずついていく"),
    "take advantage of": ("テイク アドヴァンテージ オヴ", "〜を活用する"),
    "look down on": ("ルック ダウン オン", "〜を見下す"),
    "get away with": ("ゲット アウェイ ウィズ", "罰を受けずに済む"),
    "make room for": ("メイク ルーム フォー", "〜のために場所を空ける"),
    "sort out": ("ソート アウト", "問題を解決する・整理する"),
    "put aside": ("プット アサイド", "脇に置く"),
    "take apart": ("テイク アパート", "分解する"),
    "get along with": ("ゲット アロング ウィズ", "〜とうまく付き合う"),
    "provided that": ("プロヴァイディッド ザット", "〜という条件で"),
    "now that": ("ナウ ザット", "今や〜なので"),
    "even if": ("イーヴン イフ", "たとえ〜でも"),
    "at once": ("アット ワンス", "すぐに"),
    "at least": ("アット リースト", "少なくとも"),
    "at first": ("アット ファースト", "最初は"),
    "at last": ("アット ラスト", "ついに"),
    "in case": ("イン ケイス", "〜の場合に備えて"),
    "as if": ("アズ イフ", "まるで〜のように"),
    "so that": ("ソウ ザット", "〜するために"),
}


PARAGRAPH_ENTRIES = {
    "set-01": {
        "長文語句 2A": [
            [("narrow stone channels", "ナロウ ストーン チャネルズ", "狭い石造りの水路"), ("modern water pipes", "モダン ウォーター パイプス", "近代的な水道管"), ("depended on", "ディペンディッド オン", "〜に頼っていた"), ("blocked channel", "ブロックト チャネル", "詰まった水路"), ("decorative features", "デコラティヴ フィーチャーズ", "装飾的な特徴")],
            [("shaped local work", "シェイプト ローカル ワーク", "地域の仕事のあり方に影響した"), ("rinsed fabric", "リンスト ファブリック", "布をすすいだ"), ("rainfall", "レインフォール", "降雨量"), ("turn schedules", "ターン スケジュールズ", "利用順の予定表"), ("households", "ハウスホウルズ", "世帯")],
            [("repair small cracks", "リペア スモール クラックス", "小さなひびを修理する"), ("city planners", "シティ プランナーズ", "都市計画担当者"), ("limited resource", "リミティッド リソース", "限られた資源"), ("regular cooperation", "レギュラー コオペレイション", "継続的な協力"), ("public responsibility", "パブリック リスポンサビリティ", "公共に対する責任")],
        ],
        "長文語句 2B": [
            [("rocky islands", "ロッキー アイランズ", "岩の多い島々"), ("build nests", "ビルド ネスツ", "巣を作る"), ("take off quickly", "テイク オフ クイックリー", "素早く飛び立つ"), ("nesting area", "ネスティング エリア", "営巣場所"), ("break apart", "ブレイク アパート", "崩れる")],
            [("harmless trackers", "ハームレス トラッカーズ", "鳥に害のない追跡装置"), ("recorded routes", "リコーディッド ルーツ", "移動経路を記録した"), ("followed coastlines", "フォロウド コーストラインズ", "海岸線に沿って進んだ"), ("developing", "ディヴェロッピング", "発達している途中の"), ("open water", "オウプン ウォーター", "陸から離れた水域")],
            [("warmer seas", "ウォーマー シーズ", "より暖かくなった海"), ("nesting cliffs", "ネスティング クリフス", "営巣する崖"), ("survive their first season", "サヴァイヴ ゼア ファースト シーズン", "最初の季節を生き延びる"), ("conservation workers", "コンサヴェイション ワーカーズ", "自然保護活動員"), ("limited time and money", "リミティッド タイム アンド マネー", "限られた時間と資金")],
        ],
        "メール 3A": [
            [("arranging", "アレンジング", "手配すること"), ("reservation", "レザヴェイション", "予約"), ("main galleries", "メイン ギャラリーズ", "主要展示室"), ("protect objects", "プロテクト オブジェクツ", "収蔵品を保護する"), ("basic rules", "ベイシック ルールズ", "基本的な規則")],
            [("staff entrance", "スタッフ エントランス", "職員用入口"), ("earlier notice", "アーリアー ノウティス", "以前の案内"), ("photo ID", "フォウトウ アイディー", "写真付き身分証明書"), ("lockers", "ロッカーズ", "ロッカー"), ("photography", "フォタグラフィー", "写真撮影")],
            [("reply by", "リプライ バイ", "〜までに返信する"), ("confirm", "コンファーム", "確認する"), ("borrow a small notepad", "ボロウ ア スモール ノウトパッド", "小さなメモ帳を借りる"), ("clipboards", "クリップボーズ", "クリップボード"), ("waiting list", "ウェイティング リスト", "キャンセル待ち名簿")],
        ],
        "長文内容 3B": [
            [("air conditioners", "エア コンディショナーズ", "エアコン"), ("pale stone", "ペイル ストウン", "淡い色の石"), ("reflect more sunlight", "リフレクト モア サンライト", "より多くの日光を反射する"), ("passes into", "パスィズ イントゥ", "〜の中へ伝わる"), ("comfortable", "カンフォタブル", "快適な")],
            [("mineral particles", "ミネラル パーティクルズ", "鉱物の微粒子"), ("special sheets or tiles", "スペシャル シーツ オア タイルズ", "特殊なシートや瓦"), ("fit older neighborhoods", "フィット オウルダー ネイバーフッズ", "古い街並みに調和する"), ("apartment blocks", "アパートメント ブロックス", "集合住宅"), ("bright white surface", "ブライト ホワイト サーフィス", "明るい白色の表面")],
            [("dusty areas", "ダスティー エリアズ", "ほこりの多い地域"), ("lose part of its effect", "ルーズ パート オヴ イッツ イフェクト", "効果の一部を失う"), ("heating needs", "ヒーティング ニーズ", "暖房需要"), ("turning point", "ターニング ポイント", "転機"), ("maintenance", "メインテナンス", "維持管理")],
            [("warehouses", "ウェアハウズィズ", "倉庫"), ("installed during normal repairs", "インストールド デュアリング ノーマル リペアズ", "通常の修理時に設置される"), ("air temperature", "エア テンパラチャー", "気温"), ("infrared rays", "インフラレッド レイズ", "赤外線"), ("ventilation", "ヴェンティレイション", "換気")],
        ],
    },
    "set-02": {
        "長文語句 2A": [
            [("practical problems", "プラクティカル プロブレムズ", "実際的な問題"), ("provide shade", "プロヴァイド シェイド", "日陰を作る"), ("hold rainwater", "ホウルド レインウォーター", "雨水を保持する"), ("drains", "ドレインズ", "排水口"), ("packed tightly", "パックト タイトリー", "固く締め固められた")],
            [("underground soil cells", "アンダーグラウンド ソイル セルズ", "地下の土壌セル"), ("pavement", "ペイヴメント", "舗装"), ("roots to spread", "ルーツ トゥ スプレッド", "根が広がること"), ("effectively", "イフェクティヴリー", "効果的に"), ("long-lasting shade", "ロング ラスティング シェイド", "長く続く日陰")],
            [("tree species", "トゥリー スピーシーズ", "樹種"), ("air pollution", "エア ポリューション", "大気汚染"), ("collect data", "コレクト デイタ", "データを集める"), ("use that pattern", "ユース ザット パターン", "その傾向を利用する"), ("urban forests", "アーバン フォレスツ", "都市の森林・街路樹群")],
        ],
        "長文語句 2B": [
            [("chemical dyes", "ケミカル ダイズ", "化学染料"), ("minerals", "ミネラルズ", "鉱物"), ("social position", "ソウシャル ポジション", "社会的地位"), ("differed from place to place", "ディファード フロム プレイス トゥ プレイス", "土地ごとに異なった"), ("offer clues", "オファー クルーズ", "手がかりを与える")],
            [("crush materials", "クラッシュ マテリアルズ", "材料を砕く"), ("substance", "サブスタンス", "物質"), ("mordant", "モーダント", "媒染剤"), ("fibers", "ファイバーズ", "繊維"), ("fade quickly", "フェイド クイックリー", "すぐ色あせる")],
            [("tiny traces", "タイニー トレイシズ", "微量の痕跡"), ("fragile fabrics", "フラジャイル ファブリックス", "傷みやすい布地"), ("storage", "ストレージ", "保管"), ("faded cloth", "フェイディッド クロス", "色あせた布"), ("reaching conclusions", "リーチング コンクルージョンズ", "結論に達すること")],
        ],
        "メール 3A": [
            [("waiting list", "ウェイティング リスト", "キャンセル待ち名簿"), ("participant list", "パーティシパント リスト", "参加者名簿"), ("left two seats open", "レフト トゥー シーツ オウプン", "2席の空きを生じさせた"), ("registered", "レジスタード", "登録した"), ("reserved one for you", "レザーヴド ワン フォー ユー", "あなたのために1席確保した")],
            [("details", "ディテイルズ", "詳細・注意事項"), ("side entrance", "サイド エントランス", "横の入口"), ("provide flour", "プロヴァイド フラワー", "小麦粉を用意する"), ("yeast", "イースト", "酵母・イースト"), ("apron and a container", "エイプロン アンド ア コンテイナー", "エプロンと容器")],
            [("food allergies", "フード アラジーズ", "食物アレルギー"), ("taste the bread", "テイスト ザ ブレッド", "パンを試食する"), ("serve a simple soup", "サーヴ ア シンプル スープ", "簡単なスープを出す"), ("attend", "アテンド", "参加する"), ("contact that person", "コンタクト ザット パーソン", "その人に連絡する")],
        ],
        "長文内容 3B": [
            [("harvest", "ハーヴェスト", "収穫"), ("planting season", "プランティング シーズン", "植え付け時期"), ("packets of seeds", "パケッツ オヴ シーズ", "種の小袋"), ("local varieties", "ローカル ヴァライアティーズ", "地域固有の品種"), ("grower", "グロウアー", "栽培者")],
            [("volunteers", "ヴォランティアズ", "ボランティア"), ("envelope", "エンヴェロウプ", "封筒"), ("cross-pollinated", "クロス ポリネイティッド", "交雑受粉した"), ("vine crops", "ヴァイン クロップス", "つる性作物"), ("diseased plants", "ディズィーズド プランツ", "病気にかかった植物")],
            [("commercial seeds", "コマーシャル シーズ", "市販の種"), ("convenient", "コンヴィーニエント", "便利な"), ("different climates", "ディファレント クライメッツ", "異なる気候"), ("quality control", "クオリティ コントロウル", "品質管理"), ("kept records", "ケプト レコーズ", "記録を残した")],
            [("neighborhoods", "ネイバーフッズ", "地域・近隣地区"), ("dry weeks", "ドライ ウィークス", "雨の少ない週"), ("one generation to the next", "ワン ジェネレイション トゥ ザ ネクスト", "一世代から次の世代へ"), ("traded", "トレイディッド", "交換・取引された"), ("depends on patient volunteers", "ディペンズ オン ペイシェント ヴォランティアズ", "辛抱強いボランティアに依存する")],
        ],
    },
    "set-03": {
        "長文語句 2A": [
            [("maker's marks", "メイカーズ マークス", "製造者印"), ("random scratches", "ランダム スクラッチズ", "無作為な引っかき傷"), ("brickmakers", "ブリックメイカーズ", "れんが職人"), ("were fired", "ワー ファイアド", "焼成された"), ("kiln", "キルン", "窯")],
            [("delivery", "デリヴァリー", "納品・配送品"), ("report the problem", "リポート ザ プロブレム", "問題を報告する"), ("merchants", "マーチャンツ", "商人"), ("counted goods", "カウンティッド グッズ", "商品を数えた"), ("place of production", "プレイス オヴ プロダクション", "生産地")],
            [("local historians", "ローカル ヒストリアンズ", "地域史研究者"), ("railway records", "レイルウェイ レコーズ", "鉄道記録"), ("advertisements", "アドヴァタイズメンツ", "広告"), ("routes were important", "ルーツ ワー インポータント", "どの経路が重要だったか"), ("such evidence", "サッチ エヴィデンス", "そのような証拠")],
        ],
        "長文語句 2B": [
            [("short flashes", "ショート フラッシズ", "短い発光"), ("species", "スピーシーズ", "種"), ("timing of a flash", "タイミング オヴ ア フラッシュ", "発光のタイミング"), ("background brighter", "バックグラウンド ブライター", "背景をより明るくする"), ("fewer pairs", "フューアー ペアズ", "より少ないつがい")],
            [("grassy areas", "グラッシー エリアズ", "草地"), ("reduced flashing", "リデュースト フラッシング", "発光を減らした"), ("disappear", "ディサピア", "姿を消す"), ("artificial light", "アーティフィシャル ライト", "人工照明"), ("scaring them away", "スケアリング ゼム アウェイ", "彼らを怖がらせて追い払うこと")],
            [("shielded lamps", "シールディッド ランプス", "遮光されたランプ"), ("unnecessary lights", "アンネセサリー ライツ", "不要な照明"), ("bulbs", "バルブズ", "電球"), ("less welcoming", "レス ウェルカミング", "魅力や安心感がより少ない"), ("respects both human needs", "リスペクツ ボウス ヒューマン ニーズ", "人間側の必要性も尊重する")],
        ],
        "メール 3A": [
            [("original poster", "オリジナル ポウスター", "元のポスター"), ("repairs", "リペアズ", "修理"), ("nearby walkway", "ニアバイ ウォークウェイ", "近くの歩道"), ("moved the event", "ムーヴド ジ イヴェント", "催しを移した"), ("directions", "ダイレクションズ", "道順")],
            [("indoors", "インドアズ", "屋内で"), ("water-resistant paper", "ウォーター レジスタント ペイパー", "耐水紙"), ("recycled matte paper", "リサイクルド マット ペイパー", "再生マット紙"), ("as long as", "アズ ロング アズ", "〜という条件なら"), ("colored text", "カラード テキスト", "色付きの文字")],
            [("confirm", "コンファーム", "確認する"), ("total price", "トウタル プライス", "合計金額"), ("earliest time", "アーリエスト タイム", "最も早い時刻"), ("volunteers", "ヴォランティアズ", "ボランティア"), ("delay", "ディレイ", "遅れ")],
        ],
        "長文内容 3B": [
            [("peaceful", "ピースフル", "静かで穏やかな"), ("narrow street", "ナロウ ストリート", "狭い通り"), ("sound maps", "サウンド マップス", "音の地図"), ("near airports", "ニア エアポーツ", "空港の近くで"), ("pleasant sounds", "プレザント サウンズ", "心地よい音")],
            [("planned walking route", "プランド ウォーキング ルート", "計画された歩行経路"), ("sound meter", "サウンド ミーター", "騒音計"), ("fixed points", "フィクスト ポインツ", "決められた地点"), ("loudness", "ラウドネス", "音の大きさ"), ("typical conditions", "ティピカル コンディションズ", "通常の状況")],
            [("quiet alley", "クワイエット アリー", "静かな路地"), ("measurement", "メジャメント", "測定値"), ("misleading", "ミスリーディング", "誤解を招く"), ("repeating measurements", "リピーティング メジャメンツ", "測定を繰り返すこと"), ("noise charts", "ノイズ チャーツ", "騒音図表")],
            [("practical ways", "プラクティカル ウェイズ", "実用的な使い方"), ("outdoor events", "アウトドア イヴェンツ", "屋外行事"), ("people's conversations", "ピープルズ カンヴァセイションズ", "人々の会話"), ("out of date", "アウト オヴ デイト", "情報が古くなった"), ("are combined", "アー コンバインド", "組み合わされる")],
        ],
    },
}


def load_sets():
    files_json = json.dumps(SOURCE_FILES, ensure_ascii=False)
    keys_json = json.dumps(SET_KEYS)
    node_script = f"""
const fs = require('fs');
const vm = require('vm');
const context = {{ window: {{}} }};
context.window.window = context.window;
vm.createContext(context);
for (const file of {files_json}) {{
  vm.runInContext(fs.readFileSync(file, 'utf8'), context, {{ filename: file }});
}}
const keys = new Set({keys_json});
const sets = context.window.examData.grades.grade2.sets.filter((set) => keys.has(set.key));
process.stdout.write(JSON.stringify(sets.map((set) => ({{
  key: set.key,
  label: set.label,
  readingPages: set.readingPages,
}}))));
"""
    completed = subprocess.run(
        ["node", "-e", node_script],
        cwd=ROOT,
        check=True,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )
    return json.loads(completed.stdout)


def normalize_text(text):
    return re.sub(r"[^a-z0-9]+", " ", str(text).lower().replace("’", "'")).strip()


def target_body_paragraphs(page):
    passage = page.get("passage") or []
    if page.get("label") == "メール 3A":
        return passage[2:5]
    return passage


def validate_source_and_entries(sets):
    if [item["key"] for item in sets] != list(SET_KEYS):
        raise RuntimeError(f"対象回の順序が不正です: {[item['key'] for item in sets]}")
    total_entries = 0
    for set_item in sets:
        set_key = set_item["key"]
        pages = set_item.get("readingPages") or []
        if len(pages) != 5:
            raise RuntimeError(f"{set_key}: Readingページが5ではありません")
        short_questions = pages[0].get("questions") or []
        if len(short_questions) != 17:
            raise RuntimeError(f"{set_key}: 短文語句が17問ではありません")
        for question in short_questions:
            choices = question.get("choices") or []
            if len(choices) != 4 or question.get("correct") not in (1, 2, 3, 4):
                raise RuntimeError(f"{set_key} No.{question.get('id')}: 4択または正答が不正です")
            for choice in choices:
                if choice.lower() not in CHOICE_LEXICON:
                    raise RuntimeError(f"語義辞書がありません: {set_key} No.{question.get('id')} {choice}")
        total_entries += 68

        expected_labels = ["長文語句 2A", "長文語句 2B", "メール 3A", "長文内容 3B"]
        if [page.get("label") for page in pages[1:]] != expected_labels:
            raise RuntimeError(f"{set_key}: 対象ページの順序が不正です")
        paragraph_total = 0
        for page in pages[1:]:
            label = page["label"]
            paragraphs = target_body_paragraphs(page)
            expected_count = 4 if label == "長文内容 3B" else 3
            if len(paragraphs) != expected_count:
                raise RuntimeError(f"{set_key} {label}: 本文段落数が{expected_count}ではありません")
            entry_groups = PARAGRAPH_ENTRIES[set_key][label]
            if len(entry_groups) != expected_count:
                raise RuntimeError(f"{set_key} {label}: 語彙段落数が一致しません")
            for paragraph_index, (paragraph, entries) in enumerate(zip(paragraphs, entry_groups), start=1):
                if len(entries) != 5:
                    raise RuntimeError(f"{set_key} {label} 第{paragraph_index}段落: 5語ではありません")
                normalized_paragraph = normalize_text(paragraph)
                for term, kana, meaning in entries:
                    if normalize_text(term) not in normalized_paragraph:
                        raise RuntimeError(f"本文に語がありません: {set_key} {label} 第{paragraph_index}段落 / {term}")
                    if not kana or not meaning:
                        raise RuntimeError(f"発音または意味が空です: {term}")
                paragraph_total += len(entries)
        if paragraph_total != 65:
            raise RuntimeError(f"{set_key}: 本文語彙が65ではありません: {paragraph_total}")
        total_entries += paragraph_total
    if total_entries != 399:
        raise RuntimeError(f"総項目数が399ではありません: {total_entries}")
    return total_entries


def set_run_font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Yu Gothic")


def configure_style(style, size, color, before, after, line_spacing, bold=False):
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Yu Gothic")
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line_spacing


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def add_bottom_rule(paragraph, color="C8D1DC"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = p_pr.find(qn("w:pBdr"))
    if p_borders is None:
        p_borders = OxmlElement("w:pBdr")
        p_pr.append(p_borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_borders.append(bottom)


def configure_document():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_style(document.styles["Normal"], 11, INK, 0, 6, 1.25)
    configure_style(document.styles["Heading 1"], 16, BLUE, 18, 10, 1.0, bold=True)
    configure_style(document.styles["Heading 2"], 13, BLUE, 14, 7, 1.0, bold=True)
    configure_style(document.styles["Heading 3"], 12, DARK_BLUE, 10, 5, 1.0, bold=True)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        document.styles[style_name].paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    header.paragraph_format.line_spacing = 1.0
    set_run_font(header.add_run("英検2級｜第1回〜第3回 リーディング重要語彙"), size=8.5, color=MUTED, bold=True)
    add_bottom_rule(header)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    footer.paragraph_format.line_spacing = 1.0
    set_run_font(footer.add_run("Page "), size=8.5, color=MUTED)
    add_page_field(footer)
    return document


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    tr_pr.append(table_header)


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_pr = table._tbl.tblPr
    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    table_width = table_pr.find(qn("w:tblW"))
    table_width.set(qn("w:w"), "9360")
    table_width.set(qn("w:type"), "dxa")
    indent = table_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_cell_text(cell, text, *, bold=False, color=INK, size=8.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    paragraph.paragraph_format.keep_together = True
    set_run_font(paragraph.add_run(str(text)), size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_vocab_table(document, rows):
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [1944, 2520, 2232, 2664]  # 1.35, 1.75, 1.55, 1.85 in; sum 9360 DXA.
    headers = ("出典", "英単語・熟語", "カタカナ発音", "文脈に合う日本語")
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True, color=DARK_BLUE, size=8.5)
        shade_cell(cell, HEADER_FILL)
    set_repeat_table_header(table.rows[0])
    for source, term, kana, meaning in rows:
        row = table.add_row()
        cells = row.cells
        set_cell_text(cells[0], source, color=MUTED, size=8.0)
        set_cell_text(cells[1], term, bold="✓正答" in source, size=8.5)
        set_cell_text(cells[2], kana, size=8.2)
        set_cell_text(cells[3], meaning, size=8.2)
    set_table_geometry(table, widths)
    return table


def add_cover(document):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(82)
    paragraph.paragraph_format.space_after = Pt(16)
    set_run_font(paragraph.add_run("英検2級 S-CBT TRAINING"), size=11, color=BLUE, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_run_font(title.add_run("第1回〜第3回\nリーディング重要語彙リスト"), size=25, color=INK, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    set_run_font(subtitle.add_run("短文4択を全収録｜本文は各段落5語｜全399項目"), size=12.5, color=DARK_BLUE, bold=True)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(18)
    note.paragraph_format.space_after = Pt(6)
    note.paragraph_format.line_spacing = 1.25
    set_run_font(note.add_run("出典・英単語／熟語・カタカナ発音・文脈に合う日本語を一覧化。\n4択語彙は誤答語も省略せず、正答に ✓正答 を表示しています。"), size=10.5, color=MUTED)

    count = document.add_paragraph()
    count.alignment = WD_ALIGN_PARAGRAPH.CENTER
    count.paragraph_format.space_before = Pt(26)
    count.paragraph_format.space_after = Pt(0)
    set_run_font(count.add_run("各回133項目 × 3回 ＝ 399項目"), size=12, color=BLUE, bold=True)
    document.add_page_break()


def build_short_rows(set_item):
    rows = []
    for question in set_item["readingPages"][0]["questions"]:
        for choice_index, choice in enumerate(question["choices"], start=1):
            kana, meaning = CHOICE_LEXICON[choice.lower()]
            marker = " ✓正答" if choice_index == question["correct"] else ""
            source = f"短文語句 No.{question['id']} 選択肢{choice_index}{marker}"
            rows.append((source, choice, kana, meaning))
    return rows


def build_paragraph_rows(set_item, page):
    rows = []
    for paragraph_index, entries in enumerate(PARAGRAPH_ENTRIES[set_item["key"]][page["label"]], start=1):
        for term, kana, meaning in entries:
            rows.append((f"{page['label']} 第{paragraph_index}段落", term, kana, meaning))
    return rows


def add_set_chapter(document, set_item, first=False):
    if not first:
        document.add_page_break()
    document.add_heading(set_item["label"], level=1)
    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    set_run_font(intro.add_run("短文語句68項目＋本文65項目＝133項目"), size=10, color=MUTED, bold=True)

    document.add_heading("短文語句17問 — 4択すべて", level=2)
    citation = document.add_paragraph("出典: アプリ収録の短文語句 No.1〜17。✓正答はアプリ正答データに基づきます。")
    citation.paragraph_format.space_before = Pt(4)
    citation.paragraph_format.space_after = Pt(4)
    set_run_font(citation.runs[0], size=8.5, color=MUTED)
    short_rows = build_short_rows(set_item)
    for chunk_index, chunk_start in enumerate(range(0, len(short_rows), 16)):
        if chunk_index:
            document.add_page_break()
            continuation = document.add_paragraph()
            continuation.paragraph_format.space_after = Pt(5)
            set_run_font(continuation.add_run(f"{set_item['label']}｜短文語句17問（続き）"), size=10, color=DARK_BLUE, bold=True)
        add_vocab_table(document, short_rows[chunk_start:chunk_start + 16])

    for page in set_item["readingPages"][1:]:
        document.add_page_break()
        document.add_heading(page["label"], level=2)
        citation = document.add_paragraph(f"出典: {page['passageTitle']}。本文の各対象段落から重要語を5項目ずつ選定。")
        citation.paragraph_format.space_before = Pt(4)
        citation.paragraph_format.space_after = Pt(4)
        set_run_font(citation.runs[0], size=8.5, color=MUTED)
        paragraph_rows = build_paragraph_rows(set_item, page)
        chunk_size = 10 if len(paragraph_rows) > 15 else 15
        for chunk_index, chunk_start in enumerate(range(0, len(paragraph_rows), chunk_size)):
            if chunk_index:
                document.add_page_break()
                continuation = document.add_paragraph()
                continuation.paragraph_format.space_after = Pt(5)
                set_run_font(continuation.add_run(f"{set_item['label']}｜{page['label']}（続き）"), size=10, color=DARK_BLUE, bold=True)
            add_vocab_table(document, paragraph_rows[chunk_start:chunk_start + chunk_size])


def audit_docx(document, expected_entries):
    entry_rows = 0
    for table in document.tables:
        if not table.rows or len(table.columns) != 4:
            raise RuntimeError("4列表ではない表があります")
        if [cell.text for cell in table.rows[0].cells] != ["出典", "英単語・熟語", "カタカナ発音", "文脈に合う日本語"]:
            raise RuntimeError("表ヘッダーが不正です")
        entry_rows += len(table.rows) - 1
    if entry_rows != expected_entries:
        raise RuntimeError(f"DOCX表の項目数が不正です: {entry_rows}")
    section = document.sections[0]
    if section.page_width != Inches(8.5) or section.page_height != Inches(11):
        raise RuntimeError("ページ寸法がLetterではありません")
    if any(margin != Inches(1) for margin in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin)):
        raise RuntimeError("余白が1インチではありません")


def main():
    sets = load_sets()
    total_entries = validate_source_and_entries(sets)
    document = configure_document()
    add_cover(document)
    for index, set_item in enumerate(sets):
        add_set_chapter(document, set_item, first=index == 0)
    audit_docx(document, total_entries)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(json.dumps({"output": str(OUTPUT_PATH), "sets": len(sets), "entries": total_entries}, ensure_ascii=False))


if __name__ == "__main__":
    main()
